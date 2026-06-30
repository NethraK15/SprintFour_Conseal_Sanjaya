import io
import fitz  # PyMuPDF
from docx import Document


def parse_pdf(content: bytes) -> str:
    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def parse_docx(content: bytes) -> str:
    file_stream = io.BytesIO(content)
    doc = Document(file_stream)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def parse_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore")


def parse_document(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(content)
    if lower.endswith(".docx"):
        return parse_docx(content)
    if lower.endswith(".txt"):
        return parse_txt(content)
    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")
